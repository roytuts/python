from docx import Document

document = Document()

#header section
header_section = document.sections[0]
header = header_section.header

#footer section
footer_section = document.sections[0]
footer = footer_section.footer

#header text
header_text = header.paragraphs[0]
header_text.text = "Header of my document"

#heading of the document
document.add_heading('Document Title', 0)


#paragraph
p = document.add_paragraph('A plain paragraph')

#footer text
footer_text = footer.paragraphs[0]
footer_text.text = "Footer of my document"

#write to docx
document.save('simple.docx')